#!/usr/bin/env python3
"""
Word Document Builder
Generates a beautifully formatted .docx CMA report for Hall Collins.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import io
import base64
from datetime import date
import requests


# ── Brand Colors ──────────────────────────────────────────────────────────────
NAVY = RGBColor(0x17, 0x33, 0x48)      # #173348
PINK = RGBColor(0xE9, 0x1E, 0x63)      # #E91E63
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
MID_GRAY = RGBColor(0x99, 0x99, 0x99)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_cell_border(cell, sides=("top", "bottom", "left", "right"),
                     color="173348", sz="4"):
    """Add borders to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), sz)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _para_style(para, bold=False, italic=False, size_pt=11,
                color: RGBColor = None, align=WD_ALIGN_PARAGRAPH.LEFT,
                font_name="Georgia"):
    """Apply formatting to every run in a paragraph."""
    para.alignment = align
    for run in para.runs:
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size_pt)
        run.font.name = font_name
        if color:
            run.font.color.rgb = color


def _add_heading(doc: Document, text: str, level: int = 1):
    """Add a branded section heading."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Georgia"
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = NAVY
        # Add a pink bottom border via a paragraph border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "E91E63")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = PINK
    doc.add_paragraph()  # Spacing after heading


def _add_divider(doc: Document):
    """Add a thin pink horizontal rule."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E91E63")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _fmt_val(val, prefix="", suffix="", fallback="—"):
    """Format a value safely, return fallback if None/empty."""
    if val is None or val == "" or val == 0:
        return fallback
    try:
        if isinstance(val, float) and val == int(val):
            val = int(val)
        return f"{prefix}{val:,}{suffix}" if isinstance(val, (int, float)) else f"{prefix}{val}{suffix}"
    except Exception:
        return str(val)


# ── Cover Page ────────────────────────────────────────────────────────────────

def _build_cover(doc: Document, subject: dict, logo_path: str):
    """Build a branded cover page."""
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Logo
    # Prefer the RGB-converted docx-safe version
    docx_logo = logo_path.replace('.png', '_docx.png') if '_docx' not in logo_path else logo_path
    if os.path.exists(docx_logo):
        logo_path = docx_logo
    if logo_path and os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(3.0))
    doc.add_paragraph()

    # Tagline
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tag.add_run("COMPARATIVE MARKET ANALYSIS")
    run.bold = True
    run.font.name = "Georgia"
    run.font.size = Pt(11)
    run.font.color.rgb = PINK
    run.font.all_caps = True

    doc.add_paragraph()

    # Property Address (large)
    addr = doc.add_paragraph()
    addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = addr.add_run(subject.get("street_address", "").upper())
    run.bold = True
    run.font.name = "Georgia"
    run.font.size = Pt(26)
    run.font.color.rgb = NAVY

    city_p = doc.add_paragraph()
    city_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = city_p.add_run(subject.get("city_state", ""))
    run.font.name = "Georgia"
    run.font.size = Pt(16)
    run.font.color.rgb = MID_GRAY

    doc.add_paragraph()
    doc.add_paragraph()

    # Prepared by block
    prep = doc.add_paragraph()
    prep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prep.add_run("Prepared by Hall Collins Real Estate Group")
    run.font.name = "Georgia"
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(date.today().strftime("%B %d, %Y"))
    run.font.name = "Georgia"
    run.font.size = Pt(11)
    run.font.color.rgb = MID_GRAY
    run.italic = True

    doc.add_page_break()


# ── Comparison Table ──────────────────────────────────────────────────────────

def _build_comparison_table(doc: Document, subject: dict, comps: list):
    """Build a styled comparison table: Subject vs. Comp 1/2/3."""

    _add_heading(doc, "Comparable Properties — Side by Side", level=1)

    rows_def = [
        ("Sale / List Price",   "sale_price",       "$",  ""),
        ("Bedrooms",            "beds",             "",   ""),
        ("Full Bathrooms",      "baths",            "",   ""),
        ("Garage Spaces",       "garage",           "",   ""),
        ("Living Area (sq ft)", "sqft",             "",   " sq ft"),
        ("Lot Size (acres)",    "lot_acres",        "",   " ac"),
        ("Year Built",          "year_built",       "",   ""),
        ("Days on Market",      "days_on_market",   "",   " days"),
        ("Property Type",       "property_type",    "",   ""),
        ("Quality of Finishes", "finishes_note",    "",   ""),
    ]

    col_count = 1 + min(len(comps), 3) + 1  # Feature | Subject | Comp1 | Comp2 | Comp3
    table = doc.add_table(rows=len(rows_def) + 1, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # ── Header row ──
    headers = [""] + ["⭐ Subject Property"] + [f"Comp #{i+1}" for i in range(min(len(comps), 3))]
    header_colors = ["173348", "E91E63"] + ["173348"] * 3

    for col_idx, (hdr_text, hdr_color) in enumerate(zip(headers[:col_count], header_colors[:col_count])):
        cell = table.cell(0, col_idx)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr_text)
        run.bold = True
        run.font.name = "Georgia"
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        _set_cell_bg(cell, hdr_color)

    # ── Data rows ──
    all_props = [subject] + list(comps[:3])
    for row_idx, (label, key, prefix, suffix) in enumerate(rows_def):
        real_row = row_idx + 1
        # Alternate row shading
        row_bg = "F9F4F6" if row_idx % 2 == 0 else "FFFFFF"

        # Label cell
        label_cell = table.cell(real_row, 0)
        label_cell.text = ""
        lp = label_cell.paragraphs[0]
        run = lp.add_run(label)
        run.bold = True
        run.font.name = "Georgia"
        run.font.size = Pt(9)
        run.font.color.rgb = NAVY
        _set_cell_bg(label_cell, "EEF1F4")

        # Value cells
        for col_idx, prop in enumerate(all_props[:col_count - 1]):
            val_cell = table.cell(real_row, col_idx + 1)
            val_cell.text = ""
            vp = val_cell.paragraphs[0]
            vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            raw = prop.get(key)
            display = _fmt_val(raw, prefix=prefix, suffix=suffix)
            run = vp.add_run(display)
            run.font.name = "Georgia"
            run.font.size = Pt(9)
            if col_idx == 0:  # Subject column slight highlight
                run.font.color.rgb = NAVY
                run.bold = True
            else:
                run.font.color.rgb = DARK_GRAY
            _set_cell_bg(val_cell, row_bg)

    # Set column widths
    col_widths = [Inches(1.8)] + [Inches(1.2)] * (col_count - 1)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = col_widths[min(idx, len(col_widths)-1)]

    doc.add_paragraph()

    # Comp address reference list
    p = doc.add_paragraph()
    run = p.add_run("Comparable Property Addresses:")
    run.bold = True
    run.font.name = "Georgia"
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY

    for i, comp in enumerate(comps[:3]):
        p2 = doc.add_paragraph(style="List Bullet")
        run2 = p2.add_run(f"Comp #{i+1}: {comp.get('address', 'Unknown')}  |  Source: {comp.get('source', 'Manual Entry')}")
        run2.font.name = "Georgia"
        run2.font.size = Pt(9)
        run2.font.color.rgb = DARK_GRAY
        if comp.get("url"):
            run2.font.color.rgb = RGBColor(0x17, 0x33, 0x48)

    doc.add_paragraph()


# ── Map Section ───────────────────────────────────────────────────────────────

def _build_map_section(doc: Document, subject: dict, comps: list,
                       map_png_bytes: bytes | None, anr_url: str | None):
    """Add the map section."""
    _add_heading(doc, "Location Map", level=1)

    if map_png_bytes:
        try:
            img_stream = io.BytesIO(map_png_bytes)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_stream, width=Inches(6.0))
            doc.add_paragraph()
        except Exception:
            map_png_bytes = None

    if not map_png_bytes:
        note = doc.add_paragraph()
        run = note.add_run(
            "📍 Interactive map available in the web app. "
            "Use the links below to view property locations."
        )
        run.italic = True
        run.font.name = "Georgia"
        run.font.size = Pt(10)
        run.font.color.rgb = MID_GRAY
        doc.add_paragraph()

    # ANR link
    if anr_url:
        anr_p = doc.add_paragraph()
        run = anr_p.add_run("🌿 Vermont ANR Natural Resource Atlas: ")
        run.bold = True
        run.font.name = "Georgia"
        run.font.size = Pt(10)
        run.font.color.rgb = NAVY
        run2 = anr_p.add_run(anr_url)
        run2.font.name = "Georgia"
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x17, 0x33, 0x48)
        run2.underline = True

    # Map legend
    legend_p = doc.add_paragraph()
    run = legend_p.add_run("Map Legend:  ")
    run.bold = True
    run.font.name = "Georgia"
    run.font.size = Pt(9)
    run.font.color.rgb = NAVY
    legend_items = ["★ Subject Property (Pink)"] + [f"● Comp #{i+1}" for i in range(len(comps[:3]))]
    run2 = legend_p.add_run("  |  ".join(legend_items))
    run2.font.name = "Georgia"
    run2.font.size = Pt(9)
    run2.font.color.rgb = DARK_GRAY

    doc.add_paragraph()


# ── Research Notes ────────────────────────────────────────────────────────────

def _build_research_notes(doc: Document, subject: dict, recommendations: list):
    """Build the narrative research notes section."""
    _add_heading(doc, "Research Notes & Property Overview", level=1)

    # Property narrative
    beds = subject.get("beds", "")
    baths = subject.get("baths", "")
    sqft = subject.get("sqft", "")
    year = subject.get("year_built", "")
    lot = subject.get("lot_acres", "")
    prop_type = subject.get("property_type", "home")
    features = subject.get("features_notes", "")
    description = subject.get("description", "")

    beds_str = f"{int(beds)}-bedroom" if beds else ""
    baths_str = f"{int(baths) if float(baths)==int(float(baths)) else baths}-bathroom" if baths else ""
    sqft_str = f"{int(sqft):,} square foot" if sqft else ""
    year_str = f"built in {int(year)}" if year else ""
    lot_str = f"set on {lot} acres" if lot else ""
    type_str = str(prop_type).lower().replace("_", " ") if prop_type else "home"

    parts = [p for p in [beds_str, baths_str, sqft_str] if p]
    descriptor = ", ".join(parts) if parts else ""
    location_parts = [p for p in [year_str, lot_str] if p]
    location_str = " and ".join(location_parts) if location_parts else ""

    narrative = (
        f"This lovely {descriptor} {type_str} "
        + (f"is {location_str}" if location_str else "")
        + " and offers a wonderful opportunity in today's market."
    )
    if features:
        narrative += f" {features}"
    if description:
        # Use first 2 sentences of scraped description
        sentences = description.replace("\n", " ").split(". ")
        excerpt = ". ".join(sentences[:2]).strip()
        if excerpt and len(excerpt) > 20:
            narrative += f" {excerpt}."

    p = doc.add_paragraph()
    run = p.add_run(narrative)
    run.font.name = "Georgia"
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY

    doc.add_paragraph()

    # Recommendations / Checklist
    rec_map = {
        "wait_spring": {
            "title": "🌸 Wait for Spring",
            "body": (
                "Given current market conditions and the seasonal nature of Vermont real estate, "
                "we recommend waiting for spring to list. Inventory is lower and buyer activity "
                "is significantly higher between April and June, which typically supports stronger "
                "offers and shorter days on market."
            ),
        },
        "septic_inspection": {
            "title": "🔍 Septic Inspection Recommended in Advance",
            "body": (
                "Vermont buyers frequently request septic inspections, and an unexpected failure "
                "can delay or derail a closing. We strongly recommend having the septic system "
                "professionally inspected prior to listing. A clean inspection report can be a "
                "powerful marketing tool and removes a major point of buyer uncertainty."
            ),
        },
        "home_inspection": {
            "title": "🏠 Home Inspection Recommended in Advance",
            "body": (
                "A pre-listing home inspection allows you to identify and address issues on your "
                "own timeline and budget — rather than during contract negotiations. This builds "
                "buyer confidence and can help support your asking price."
            ),
        },
        "staging": {
            "title": "🛋️ Staging Instructions",
            "body": (
                "First impressions are everything. We recommend decluttering all living spaces, "
                "depersonalizing the home (removing family photos, personal collections), "
                "ensuring all rooms have adequate lighting, and adding fresh flowers or plants "
                "to key areas. Consider a professional stager consultation for the main living areas."
            ),
        },
        "deep_clean": {
            "title": "🧹 Deep Clean & Clear Out Recommended",
            "body": (
                "We recommend a professional deep clean prior to listing photography and showings. "
                "Pay particular attention to kitchens, bathrooms, windows, and floors. "
                "Additionally, clearing out attics, basements, and garages signals to buyers "
                "that the home has been well cared for and makes spaces appear larger."
            ),
        },
        "land_subdivision": {
            "title": "📐 Land Subdivision Opportunity",
            "body": (
                "The lot size and configuration may present an opportunity for subdivision, "
                "which could significantly increase the overall value of the property. "
                "We recommend consulting with a local surveyor and the town planning department "
                "to explore this option before listing."
            ),
        },
        "painting_projects": {
            "title": "�� Painting / Complete A Few Projects",
            "body": (
                "A fresh coat of neutral paint is one of the highest-return investments before "
                "listing. We recommend addressing any visible deferred maintenance — such as "
                "peeling paint, cracked trim, or incomplete renovations — prior to going to "
                "market. Buyers tend to over-discount for visible cosmetic issues."
            ),
        },
    }

    if recommendations:
        _add_heading(doc, "Agent Recommendations", level=2)
        for rec_key in recommendations:
            if rec_key in rec_map:
                rec = rec_map[rec_key]
                title_p = doc.add_paragraph()
                run = title_p.add_run(rec["title"])
                run.bold = True
                run.font.name = "Georgia"
                run.font.size = Pt(11)
                run.font.color.rgb = NAVY

                body_p = doc.add_paragraph()
                run2 = body_p.add_run(rec["body"])
                run2.font.name = "Georgia"
                run2.font.size = Pt(10)
                run2.font.color.rgb = DARK_GRAY
                body_p.paragraph_format.left_indent = Inches(0.3)
                doc.add_paragraph()

    doc.add_paragraph()


# ── Price Recommendation ──────────────────────────────────────────────────────

def _build_price_recommendation(doc: Document, subject: dict,
                                 comps: list, price_low: int, price_high: int,
                                 price_notes: str):
    """Build the price recommendation section."""
    _add_heading(doc, "Price Recommendation", level=1)

    # Price range box (simulated with a shaded table)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    labels = ["Recommended Low", "Target List Price", "Recommended High"]
    prices = [
        price_low,
        int((price_low + price_high) / 2),
        price_high,
    ]
    bg_colors = ["EEF1F4", "173348", "EEF1F4"]
    txt_colors = [NAVY, WHITE, NAVY]

    for i, (lbl, price, bg, txt_col) in enumerate(zip(labels, prices, bg_colors, txt_colors)):
        cell = table.cell(0, i)
        cell.text = ""
        _set_cell_bg(cell, bg)
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p1.add_run(lbl + "\n")
        run.font.name = "Georgia"
        run.font.size = Pt(9)
        run.font.color.rgb = txt_col
        run2 = p1.add_run(f"${price:,}")
        run2.bold = True
        run2.font.name = "Georgia"
        run2.font.size = Pt(16)
        run2.font.color.rgb = PINK if i == 1 else txt_col
        cell.width = Inches(2.0)

    doc.add_paragraph()
    doc.add_paragraph()

    # Comp price analysis
    comp_prices = [c.get("sale_price") for c in comps[:3] if c.get("sale_price")]
    if comp_prices:
        avg_comp = int(sum(comp_prices) / len(comp_prices))
        low_comp = min(comp_prices)
        high_comp = max(comp_prices)

        analysis = doc.add_paragraph()
        run = analysis.add_run("Comparable Sales Summary:  ")
        run.bold = True
        run.font.name = "Georgia"
        run.font.size = Pt(10)
        run.font.color.rgb = NAVY
        run2 = analysis.add_run(
            f"Based on {len(comp_prices)} comparable sale(s), "
            f"the range was ${low_comp:,} – ${high_comp:,} "
            f"with an average of ${avg_comp:,}."
        )
        run2.font.name = "Georgia"
        run2.font.size = Pt(10)
        run2.font.color.rgb = DARK_GRAY
        doc.add_paragraph()

    # Price per sq ft analysis
    subj_sqft = subject.get("sqft")
    if subj_sqft and comp_prices:
        comp_sqfts = [c.get("sqft") for c in comps[:3] if c.get("sqft") and c.get("sale_price")]
        if comp_sqfts:
            ppsf_list = [c["sale_price"] / c["sqft"] for c in comps[:3]
                         if c.get("sqft") and c.get("sale_price")]
            avg_ppsf = sum(ppsf_list) / len(ppsf_list)
            implied_price = int(avg_ppsf * subj_sqft)
            ppsf_p = doc.add_paragraph()
            run = ppsf_p.add_run("Price Per Square Foot Analysis:  ")
            run.bold = True
            run.font.name = "Georgia"
            run.font.size = Pt(10)
            run.font.color.rgb = NAVY
            run2 = ppsf_p.add_run(
                f"Average comp price/sq ft: ${avg_ppsf:.0f}/sq ft. "
                f"Applied to subject's {int(subj_sqft):,} sq ft, "
                f"this implies a value of approximately ${implied_price:,}."
            )
            run2.font.name = "Georgia"
            run2.font.size = Pt(10)
            run2.font.color.rgb = DARK_GRAY
            doc.add_paragraph()

    # Agent notes
    if price_notes:
        notes_p = doc.add_paragraph()
        run = notes_p.add_run("Agent Pricing Notes:  ")
        run.bold = True
        run.font.name = "Georgia"
        run.font.size = Pt(10)
        run.font.color.rgb = NAVY
        run2 = notes_p.add_run(price_notes)
        run2.font.name = "Georgia"
        run2.font.size = Pt(10)
        run2.font.color.rgb = DARK_GRAY
        run2.italic = True

    doc.add_paragraph()


# ── Footer ────────────────────────────────────────────────────────────────────

def _build_footer(doc: Document):
    """Add a branded footer to all pages."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"Hall Collins Real Estate Group  |  Comparative Market Analysis  |  "
        f"Confidential  |  {date.today().strftime('%B %Y')}"
    )
    run.font.name = "Georgia"
    run.font.size = Pt(8)
    run.font.color.rgb = MID_GRAY


# ── Master Build Function ─────────────────────────────────────────────────────

def build_cma_document(
    subject: dict,
    comps: list,
    recommendations: list,
    price_low: int,
    price_high: int,
    price_notes: str,
    map_png_bytes: bytes | None,
    anr_url: str | None,
    logo_path: str = "hall_collins_logo_docx.png",
) -> io.BytesIO:
    """
    Build the complete CMA Word document and return as BytesIO.

    subject: dict with keys: street_address, city_state, beds, baths, sqft,
             lot_acres, year_built, sale_price, garage, property_type,
             features_notes, description, source, finishes_note
    comps: list of dicts with same keys + address
    recommendations: list of rec_keys (strings)
    price_low, price_high: integers
    price_notes: free text agent notes on pricing
    map_png_bytes: bytes from screenshot or None
    anr_url: direct link to VT ANR Atlas
    logo_path: path to hall_collins_logo.png
    """
    doc = Document()

    # Set default font for document
    from docx.oxml.ns import qn as _qn
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Georgia"
    font.size = Pt(11)

    # Cover
    _build_cover(doc, subject, logo_path)

    # Map
    _build_map_section(doc, subject, comps, map_png_bytes, anr_url)
    _add_divider(doc)
    doc.add_paragraph()

    # Comparison Table
    _build_comparison_table(doc, subject, comps)
    _add_divider(doc)
    doc.add_paragraph()

    # Research Notes
    _build_research_notes(doc, subject, recommendations)
    _add_divider(doc)
    doc.add_paragraph()

    # Price Recommendation
    _build_price_recommendation(doc, subject, comps, price_low, price_high, price_notes)

    # Footer
    _build_footer(doc)

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
